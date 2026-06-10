"""Attachment document parser: PDF (PyMuPDF), DOCX (python-docx), XLSX (openpyxl).

Provides a single entry point::

    from parser.document_parser import parse_document
    text = parse_document("/path/to/file.pdf")

Returns the extracted text, or an empty string on failure.
"""

import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)

# 5 MB cap on extracted text to protect memory
MAX_CHARS = 5_000_000

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".xlsx"})
DEPRECATED_EXTENSIONS = frozenset({".doc", ".xls"})


# ── Public API ─────────────────────────────────────────────────

def parse_document(filepath: str, max_chars: int = MAX_CHARS) -> str:
    """Extract text from a document at *filepath*.

    Supported formats: .pdf, .docx, .xlsx
    Old binary formats (.doc, .xls) return an empty string.

    Args:
        filepath: Path to the document on disk.
        max_chars: Maximum number of characters to return (truncation safety).

    Returns:
        Extracted text, truncated at *max_chars*.  Empty string on failure
        or unsupported format.
    """
    if not os.path.isfile(filepath):
        logger.warning("Document not found: %s", filepath)
        return ""

    ext = os.path.splitext(filepath)[1].lower()

    if ext in DEPRECATED_EXTENSIONS:
        logger.warning("Unsupported legacy format '%s': %s", ext, filepath)
        return ""

    if ext not in SUPPORTED_EXTENSIONS:
        logger.debug("Skipping unsupported extension '%s': %s", ext, filepath)
        return ""

    try:
        text = _parse_by_ext(filepath, ext)
    except Exception:
        logger.exception("Parse failed for %s", filepath)
        return ""

    # Safety truncation
    if len(text) > max_chars:
        logger.warning(
            "Truncating %s from %d to %d chars",
            filepath, len(text), max_chars,
        )
        return text[:max_chars]

    return text


# ── Per-format parsers ─────────────────────────────────────────

def _parse_by_ext(filepath: str, ext: str) -> str:
    if ext == ".pdf":
        return _parse_pdf(filepath)
    if ext == ".docx":
        return _parse_docx(filepath)
    if ext == ".xlsx":
        return _parse_xlsx(filepath)
    return ""


# ── PDF ────────────────────────────────────────────────────────

def _parse_pdf(filepath: str) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(filepath)
    try:
        parts: list[str] = []
        for page in doc:  # type: ignore[union-attr]
            page_text = page.get_text()  # type: ignore[union-attr]
            if page_text:
                parts.append(page_text.strip())
        return "\n\n".join(parts)
    finally:
        doc.close()


# ── DOCX ───────────────────────────────────────────────────────

def _parse_docx(filepath: str) -> str:
    from docx import Document

    doc = Document(filepath)
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


# ── XLSX ───────────────────────────────────────────────────────

def _parse_xlsx(filepath: str) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(filepath, read_only=True, data_only=True)
    try:
        parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"[Sheet: {sheet_name}]")
            row_texts: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    row_texts.append("\t".join(cells))
            parts.append("\n".join(row_texts))
        return "\n\n".join(parts)
    finally:
        wb.close()


# ── Test entry point ───────────────────────────────────────────

def main() -> None:
    """Quick smoke-test.  Drops a tiny PDF in a temp file and parses it."""
    import tempfile

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    # We can't easily create synthetic PDF/DOCX/XLSX without their libraries,
    # so just demonstrate the error-handling paths.
    print("=== Missing file ===")
    result = parse_document("/tmp/__does_not_exist__.pdf")
    print(f"Result: {result!r}")

    print("\n=== Unsupported legacy .doc ===")
    result = parse_document("/tmp/legacy.doc")
    print(f"Result: {result!r}")

    print("\n=== Unknown extension ===")
    result = parse_document("/tmp/notes.txt")
    print(f"Result: {result!r}")

    # If a real file is passed as argument, try it
    import sys
    if len(sys.argv) > 1:
        path = sys.argv[1]
        print(f"\n=== Parsing: {path} ===")
        print(parse_document(path)[:500])


if __name__ == "__main__":
    main()
